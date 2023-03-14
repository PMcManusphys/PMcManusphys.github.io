

def main(fnjs):
    
    import pandas as pd
    import numpy as np
    from operator import itemgetter
    import math
    import docx
    from docx.shared import Inches, Cm
    import json
    #import requests

    from comments import concept_com, content_com, overall_com

    from weighting import time_based_conc, average_conc, average_cont, amount_conc


    '''
    #get the json file
    api_url = "https://jsonplaceholder.typicode.com/todos/1"
    response = requests.get(api_url)
    data = response.json()
    '''


    #reading json file.
    import json
    f = open(fnjs)
    data = json.load(f)
    filename = data['input_name']
    filename_ouput = data['output_name']
    just_bw = data['just_bw'] # just best and worst comments
    no_zero = data['no_zero']# if true zeros will be skipped.
    inp = data['weighting'] # this is a dictionary
    type = data['type']
    scale = float(data['scale'])
    new_old = data['new_old']
    weighting_type = data['weighting_type']
    scale_grades = data['prof_scale']
    scale_grades  = sorted(scale_grades)
    course = data['course_name']
    district = data['district']





    #here we are going to load in the data
    data = pd.read_excel(filename, header = None)
    data = np.array(data)
    #replace nan with 0
    for i,x in enumerate(data):
        for j,y in enumerate(x):
            if y != y:
                data[i][j] = 0

    #lets start by reading some of the data, first we will find all the different concepts and content
    skills = [[],[]]  #skills[0] concepts, skills[1] content
    for i in range(2):
        conc = data[i,2:]
        conc = [x for x in conc if x==x]
        
        try:
            conc = ','.join(conc)
        
        except:
            t = 10
        conc = conc.split(',')
        for idx, x in enumerate(conc):
            conc[idx] = x.rstrip()
            conc[idx] = conc[idx].lstrip()
        skills[i].append(list(set(conc)))




    assessment = []
    scaling = []
    ############### weighting ##################
    for x in inp:
        assessment.append(x[0])
        scaling.append(x[1])


    # we are going to replace the name of the assessment with the weighting
    for i, x in enumerate(data[2,2:]):
        for j, y in enumerate(assessment):
            x = str(x).rstrip()
            if x == y:
                data[2,i+2] = scaling[j]




    # now lets calculate the total grade

    grades = []
    grades_percent = []
    #type = 'proficiency scale' #change able
    #type = 'grade based'
    if type ==  'proficiency scale':
        for i in range(len(data)-5):
            marks = 0
            total = 0
            for j in range(len(data[0])-2):
                if district.lower() == "sd36":
                    if str(data[i+5,j+2]).lower() == 'ie':
                        data[i+5,j+2] = 0
                    elif str(data[i+5,j+2]).lower() == 'emg': # 'emg, dev, pro, ext '
                        data[i+5,j+2] = 1
                    elif str(data[i+5,j+2]).lower() == 'dev':    # 'emg, dev, pro, ext '
                        data[i+5,j+2] = 2
                    elif str(data[i+5,j+2]).lower() == 'pro':    # 'emg, dev, pro, ext '
                        data[i+5,j+2] = 3
                    elif str(data[i+5,j+2]).lower() == 'ext':    # 'emg, dev, pro, ext '
                        data[i+5,j+2] = 4

                    if data[i+5,j+2] == 'exc' or data[i+5,j+2] == 'exu' or data[i+5,j+2] == 'excused':
                        if no_zero and float(data[i+5,j+2]) == 0:
                            zx = 0
                if district.lower() == "sd39":  ################# fix
                    if str(data[i+5,j+2]).lower() == 'ie':
                        data[i+5,j+2] = 0
                    elif str(data[i+5,j+2]).lower() == 'beg': # 'emg, dev, pro, ext '
                        data[i+5,j+2] = 1
                    elif str(data[i+5,j+2]).lower() == 'dev':    # 'emg, dev, pro, ext '
                        data[i+5,j+2] = 2
                    elif str(data[i+5,j+2]).lower() == 'app':    # 'emg, dev, pro, ext '
                        data[i+5,j+2] = 3
                    elif str(data[i+5,j+2]).lower() == 'ext':    # 'emg, dev, pro, ext '
                        data[i+5,j+2] = 4
                    
                    if data[i+5,j+2] == 'exc' or data[i+5,j+2] == 'exu' or data[i+5,j+2] == 'excused':
                        if no_zero and float(data[i+5,j+2]) == 0:
                            zx = 0

                marks += float(data[i+5,j+2])*float(data[2,j+2])
                total += 4*float(data[2,j+2])


            if district.lower() == "sd36":
                if marks == marks: # do not if NaN
                    if round(100*marks/(total+0.001),1) < scale_grades[0]:
                        grades.append('Incomplete')
                    elif round(100*marks/(total+0.001),1) >= scale_grades[0] and round(100*marks/(total+0.001),1) < scale_grades[1]:
                        grades.append('Emerging')
                    elif round(100*marks/(total+0.001),1) >= scale_grades[1] and round(100*marks/(total+0.001),1) < scale_grades[2]:
                        grades.append('Developing')
                    elif round(100*marks/(total+0.001),1) >= scale_grades[2] and round(100*marks/(total+0.001),1) < scale_grades[3]:
                        grades.append('Proficient')
                    elif round(100*marks/(total+0.001),1) >= scale_grades[3]:
                        grades.append('Extending')

                grades_percent.append(round(100*marks/(total+0.001),1))
            if district.lower() == "sd39":  ################# fix
                if marks == marks: # do not if NaN
                    if round(100*marks/(total+0.001),1) < scale_grades[0]:
                        grades.append('Insufficient')
                    elif round(100*marks/(total+0.001),1) >= scale_grades[0] and round(100*marks/(total+0.001),1) < scale_grades[1]:
                        grades.append('Beginning')
                    elif round(100*marks/(total+0.001),1) >= scale_grades[1] and round(100*marks/(total+0.001),1) < scale_grades[2]:
                        grades.append('Developing')
                    elif round(100*marks/(total+0.001),1) >= scale_grades[2] and round(100*marks/(total+0.001),1) < scale_grades[3]:
                        grades.append('Applying')
                    elif round(100*marks/(total+0.001),1) >= scale_grades[3]:
                        grades.append('Extending')
                
                    grades_percent.append(round(100*marks/(total+0.001),1))


    else:
        #I the column j is the row, The students start at 5 and the marks start at 2 (index not number)
        for i in range(len(data)-5):
            marks = 0
            total = 0
            for j in range(len(data[0])-2):
                if not isinstance(data[i+5,j+2],str):  # if excused do not include
                    
             
                        if data[i+5,j+2] == 'exc' or data[i+5,j+2] == 'exu' or data[i+5,j+2] == 'excused':
                            if no_zero and float(data[i+5,j+2]) == 0:
                                zx = 0
                        else:
                            marks += float(data[i+5,j+2])*float(data[2,j+2])
                            total += float(data[4,j+2])*float(data[2,j+2])
        
            if marks == marks: # do not if NaN
                grades.append(round(100*marks/(total+0.001),1))   # round to 1 decimal
                grades_percent.append(round(100*marks/(total+0.001),1))

    # now lets calculate the grade for each concept
    # For concepts I have a weighting towards more recent work
    # I am going to have a sliding weighting scale for.
    # If there is 5 occurances then the weighting will be (0.5,0.6,0.7,0.8,0.9)   Feel free to change this

    #first lets find out how many of each skill
    num_skill = []
    for x in skills[0][0][0:]:
        #lets find all assignments with this concept
        numb = 0
        for j, y in enumerate(data[0,2:]): # This is the assiggnment

            if x in y: # then our concept is in this assignment
                numb += 1
        num_skill.append(numb)

    #now lets make our weighting




    #weighting_type =

    if weighting_type == "time_based":
        concs_grade = time_based_conc(num_skill,data,skills,scale,new_old) #change able
    elif weighting_type == "average":
        concs_grade = average_conc(data,skills) #change able
    elif weighting_type == "time_amount":
        concs_grade = time_amount_conc(num_skill,data,skills,scale,new_old) #change able
    elif weighting_type == "amount":
        concs_grade = amount_conc(data,skills,scale) #change able
    else:
        concs_grade = average_conc(data,skills) #change able
    # now lets calculate the grade for each content
    # no special weighting here
    conts_grade = average_cont(data, skills)

    # now we can display the data
    grade_book = [['name','total grade','comp grade','cont grade']]
    if type ==  'proficiency scale':
        for i in range(len(grades)): # I know this can be done easier I just want to visalize will replace anyways
            temp = []
            temp.append(data[i+5,1]) # name
            temp.append(grades[i]) # total grade
            c = []
            d = []
            if district.lower() == "sd36":
                for idx, j in enumerate(concs_grade[i]):
                    if j < scale_grades[0]:
                        x = 'Insufficient'
                    elif j >= scale_grades[0] and j < scale_grades[1]:
                        x = 'Emerging'
                    elif j >= scale_grades[1] and j < scale_grades[2]:
                        x = 'Developing'
                    elif j >= scale_grades[2] and j < scale_grades[3]:
                        x = 'Proficient'
                    elif j >= scale_grades[3]:
                        x = 'Extending'
                    
                    c.append(x)
            
                for idx,j in enumerate(conts_grade[i]):
       
                    if j < scale_grades[0]:
                        x = 'Insufficient'
                    elif j >= scale_grades[0] and j < scale_grades[1]:
                        x = 'Emerging'
                    elif j >= scale_grades[1] and j < scale_grades[2]:
                        x = 'Developing'
                    elif j >= scale_grades[2] and j < scale_grades[3]:
                        x = 'Proficient'
                    elif j >= scale_grades[3]:
                        x = 'Extending'
                    d.append(x)

            if district.lower() == "sd39":   ########### fix wording
                for idx, j in enumerate(concs_grade[i]):
        
                    if j < scale_grades[0]:
                        x = 'Insufficient'
                    elif j >= scale_grades[0] and j < scale_grades[1]:
                        x = 'Beginning'
                    elif j >= scale_grades[1] and j < scale_grades[2]:
                        x = 'Developing'
                    elif j >= scale_grades[2] and j < scale_grades[3]:
                        x = 'Applying'
                    elif j >= scale_grades[3]:
                        x = 'Extending'
                    c.append(x)

                for idx,j in enumerate(conts_grade[i]):
       
                    if j < scale_grades[0]:
                        x = 'Insufficient'
                    if j >= scale_grades[0] and j < scale_grades[1]:
                        x = 'Beginning'
                    if j >= scale_grades[1] and j < scale_grades[2]:
                        x = 'Developing'
                    if j >= scale_grades[2] and j < scale_grades[3]:
                        x = 'Applying'
                    if j >= scale_grades[3]:
                        x = 'Extending'
                    d.append(x)

            temp.append(c)
            temp.append(d)
            grade_book.append(temp)

    else:
        for i in range(len(grades)): # I know this can be done easier I just want to visalize will replace anyways
            temp = []
            temp.append(data[i+5,1]) # name
            temp.append(grades[i]) # total grade
            for idx, j in enumerate(concs_grade[i]):
                temp.append(skills[0][0][idx])
                temp.append(j)
            for idx,j in enumerate(conts_grade[i]):
                temp.append(skills[1][0][idx])
                temp.append(j)

            grade_book.append(temp)




    doc = docx.Document()
    sections = doc.sections
    margin = 0.65
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)

    from docx.shared import Pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(14)


    #concs_grade
    #skills[0][0][0:]


    for i in range(len(grades)): # I know this can be done easier I just want to visalize will replace anyways
            para = doc.add_paragraph()
            para.add_run(f" \n {data[i+5,1]} ").bold = True
         
            font.size = Pt(10)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = Inches(3.75)
            table.columns[1].width = Inches(4.4)


            row = table.rows[0].cells

            if type !=  'proficiency scale':
                comment = f"Overall Mark: {grade_book[i+1][1]}% \n \n"
            else:
                comment = f"Overall Mark: {grade_book[i+1][1]} \n \n"

    #competencies
            comment += "CURRICULAR COMPETENCIES: \n \n"
            temp = skills[0][0][0:]

            for  n,comp in enumerate(temp):
                if type !=  'proficiency scale':
                    comment += f"{comp.capitalize()}: {grade_book[i+1][3+2*n]}% \n \n"
                else:
                    comment += f"{comp.capitalize()}: {grade_book[i+1][2][n]} \n \n"


            comment += "CONTENT: \n \n"
            temp = skills[1][0][0:]
            x = len(skills[0][0][0:])*2 + 3
            for  n,cont in enumerate(temp):
                if type !=  'proficiency scale':
                    comment += f"{cont.capitalize()}: {grade_book[i+1][x + 2*n]}% \n \n"
                else:
                    comment += f"{cont.capitalize()}: {grade_book[i+1][3][n]} \n \n"


            row[0].text = comment


            consistent = 0
            improving = 0

            comment = overall_com(grades_percent[i], data[i+5,1], course, scale_grades,consistent,improving) # grade, name, course name
            comment += " \n"

            for idx, j in enumerate(concs_grade[i]):
                if idx < 1:
                    if j == max(concs_grade[i]):
                        bw = 1
                    elif j == min(concs_grade[i]):
                        bw = -1
                    else:
                        bw = 0
                    if not just_bw:
                        comment += concept_com(j, data[i+5,1], skills[0][0][idx],0,bw,scale_grades) #grade, name, concept
                    elif just_bw and bw != 0:
                        comment += concept_com(j, data[i+5,1], skills[0][0][idx],0,bw,scale_grades) #grade, name, concept
                else:
                    if j == max(concs_grade[i]):
                        bw = 1
                    elif j == min(concs_grade[i]):
                        bw = -1
                    else:
                        bw = 0
                    if not just_bw:
                        comment += concept_com(j, data[i+5,1], skills[0][0][idx],1,bw,scale_grades) #grade, name, concept
                    elif just_bw and bw != 0:
                        comment += concept_com(j, data[i+5,1], skills[0][0][idx],1,bw,scale_grades) #grade, name, concept

            comment += " \n \n"
            

            for idx,j in enumerate(conts_grade[i]):
                if idx < 1:
                    if j == max(conts_grade[i]):
                        bw = 1
                    elif j == min(conts_grade[i]):
                        bw = -1
                    else:
                        bw = 0
                    if not just_bw:
                        comment += content_com(j, data[i+5,1], skills[1][0][idx],0,bw, scale_grades) #grade, name, content
                    elif just_bw and bw != 0:
                        comment += content_com(j, data[i+5,1], skills[1][0][idx],0,bw, scale_grades) #grade, name, content

                else:
                    if j == max(conts_grade[i]):
                        bw = 1
                    elif j == min(conts_grade[i]):
                        bw = -1
                    else:
                        bw = 0
                    if not just_bw:
                        comment += content_com(j, data[i+5,1], skills[1][0][idx],1,bw, scale_grades) #grade, name, content
                    elif just_bw and bw != 0:
                        comment += content_com(j, data[i+5,1], skills[1][0][idx],1,bw, scale_grades) #grade, name, content

            row[1].text = comment

    doc.save(f"{filename_ouput}.docx")
    #print("Your file was created")

    #This will remove the gradebook
    #os.remove(filename)

    '''
    #sending the file to the api
    with open(f"{filename_ouput}.docx", 'rb') as f:
        r = requests.post('http://httpbin.org/post', files={f"{filename_ouput}.docx": f})  # might need to use api_url or api_url+'/post'


    '''





